import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Mm


url = input("paste an url of habr article (ex if you just want to see an example):  ")
if url == 'ex':
    url = 'https://habr.com/ru/articles/803769/'

doc_name = input("input word file name: ")
doc = Document()

img_counter = 1

response = requests.get(url)
soup = BeautifulSoup(response.text, 'lxml')
filling = soup.find('div', class_="article-formatted-body").find()

head = doc.add_heading(soup.h1.text)

for tag in filling:

    if tag.name == 'p':
        doc.add_paragraph(tag.text)

    elif tag.name == 'h2':
        doc.add_heading(tag.text, level=2)

    elif tag.name == 'h3':
        doc.add_heading(tag.text, level=3)

    elif tag.name == 'figure':
        picture = tag.find('img')

        if picture:
            #   download images because docx can`t insert files from the internet
            picture_url = picture.get('data-src')
            picture_path = f'./sources/img{img_counter}.jpg'
            img_counter += 1
            resp = requests.get(picture_url)

            with open(picture_path, 'wb') as file:
                file.write(resp.content)

            doc.add_picture(picture_path, width=Mm(140))


doc.save(f'./generated_word_files/{doc_name}.docx')
